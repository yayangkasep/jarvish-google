import os
import io


def extract_text_from_file(file_path: str, max_pages: int = 20) -> str:
    """
    Extracts text from a given document file based on its extension.
    Limits extraction to max_pages to avoid token overflow.
    Supported extensions: .pdf, .docx, .xlsx, .txt, .csv
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(file_path, max_pages)
    elif ext == ".docx":
        return _extract_docx(file_path, max_pages)
    elif ext == ".xlsx":
        return _extract_xlsx(file_path, max_pages)
    elif ext in [".txt", ".csv", ".md"]:
        return _extract_text(file_path, max_pages)
    else:
        return f"[Unsupported file format: {ext}]"


def _extract_pdf(file_path: str, max_pages: int) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        text = []
        num_pages = min(len(reader.pages), max_pages)
        for i in range(num_pages):
            page_text = reader.pages[i].extract_text()
            if page_text:
                text.append(f"--- Page {i + 1} ---\n{page_text.strip()}")
        if len(reader.pages) > max_pages:
            text.append(
                f"\n[Note: Document truncated to first {max_pages} pages for context limits.]"
            )
        return "\n\n".join(text)
    except Exception as e:
        return f"[Error reading PDF: {e}]"


def _extract_docx(file_path: str, max_pages: int) -> str:
    try:
        from docx import Document

        doc = Document(file_path)
        text = []
        # Estimate 'pages' loosely by paragraphs, ~40 paragraphs per page
        max_paragraphs = max_pages * 40
        num_paragraphs = min(len(doc.paragraphs), max_paragraphs)
        for i in range(num_paragraphs):
            p = doc.paragraphs[i].text.strip()
            if p:
                text.append(p)
        if len(doc.paragraphs) > max_paragraphs:
            text.append(f"\n[Note: Document truncated for context limits.]")
        return "\n".join(text)
    except Exception as e:
        return f"[Error reading DOCX: {e}]"


def _extract_xlsx(file_path: str, max_pages: int) -> str:
    try:
        from openpyxl import load_workbook

        wb = load_workbook(file_path, data_only=True)
        text = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text.append(f"--- Sheet: {sheet_name} ---")
            row_count = 0
            for row in sheet.iter_rows(values_only=True):
                # Filter out empty rows
                if not any(row):
                    continue
                # Format row as CSV-like
                row_str = " | ".join([str(c) if c is not None else "" for c in row])
                text.append(row_str)
                row_count += 1
                if row_count > (max_pages * 50):  # ~50 rows per 'page' limit per sheet
                    text.append("[Sheet truncated]")
                    break
        return "\n".join(text)
    except Exception as e:
        return f"[Error reading XLSX: {e}]"


def _extract_text(file_path: str, max_pages: int) -> str:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            lines = []
            max_lines = max_pages * 50
            for i, line in enumerate(f):
                if i >= max_lines:
                    lines.append("\n[Note: Document truncated for context limits.]")
                    break
                lines.append(line.strip())
            return "\n".join(lines)
    except Exception as e:
        return f"[Error reading text file: {e}]"
